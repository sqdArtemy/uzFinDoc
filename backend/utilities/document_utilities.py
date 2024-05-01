"""
This file will contain all data extractors for different datatypes which are used in this project
"""
import asyncio
import aiofiles
import PyPDF2
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document
from concurrent.futures import ThreadPoolExecutor

from utilities.functions import is_file_exists


async def get_text_from_pdf(file_path: str) -> str:
    if is_file_exists(file_path):

        async with aiofiles.open(file_path, 'rb') as pdf_file:
            pdf_bytes = await pdf_file.read()

        def process_pdf(pdf_bytes):
            pdf_stream = BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            document_text = ""
            for page_number in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_number]
                document_text += page.extract_text()
            return document_text

        with ThreadPoolExecutor() as executor:
            text = await asyncio.get_running_loop().run_in_executor(executor, process_pdf, pdf_bytes)

        return text
    else:
        raise FileNotFoundError(f"File {file_path} does not exist")


async def get_text_from_docx(file_path: str) -> str:
    if is_file_exists(file_path):
        async with aiofiles.open(file_path, 'rb') as docx_file:
            docx_bytes = await docx_file.read()

        def process_docx(docx_bytes):
            docx_stream = BytesIO(docx_bytes)
            docx_reader = Document(docx_stream)
            document_text = ""
            for paragraph in docx_reader.paragraphs:
                document_text += paragraph.text + "\n"
            return document_text

        with ThreadPoolExecutor() as executor:
            text = await asyncio.get_running_loop().run_in_executor(executor, process_docx, docx_bytes)

        return text
    else:
        raise FileNotFoundError(f"File {file_path} does not exist")


async def create_pdf_with_text(text: str, file_path: str) -> None:
    def create_pdf(text, file_path):
        c = canvas.Canvas(file_path)
        max_width = 450
        font_name = "Helvetica"
        font_size = 12
        leading = font_size * 1.2
        x_position = 100
        y_position = 750
        c.setFont(font_name, font_size)

        words = text.split()
        line = ''
        for word in words:
            if c.stringWidth(f"{line} {word}", font_name, font_size) <= max_width:
                line += f"{word} "
            else:
                c.drawString(x_position, y_position, line)
                line = f"{word} "
                y_position -= leading

        if line:
            c.drawString(x_position, y_position, line)
        c.save()

    with ThreadPoolExecutor() as executor:
        await asyncio.get_running_loop().run_in_executor(executor, create_pdf, text, file_path)


async def create_docx_with_text(text: str, file_path: str) -> None:
    def create_docx(text, file_path):
        doc = Document()
        doc.add_paragraph(text)
        doc.save(file_path)

    with ThreadPoolExecutor() as executor:
        await asyncio.get_running_loop().run_in_executor(executor, create_docx, text, file_path)
